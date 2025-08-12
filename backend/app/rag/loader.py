import os
import tempfile
import hashlib
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
from unstructured.partition.auto import partition
from langchain.schema import Document
import logging
import pandas as pd
import csv
from io import StringIO
from datetime import datetime

# Optional fallback imports
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    pdfplumber = None

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    DocxDocument = None

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Enhanced document loader with fallback processing and content hashing."""
    
    def __init__(self):
        """Initialize document loader with enhanced processing capabilities."""
        self.supported_formats = {
            "application/pdf": self._load_pdf,
            "text/plain": self._load_text,
            "text/markdown": self._load_text,
            "text/csv": self._load_csv,
            "application/csv": self._load_csv,
            "application/msword": self._load_word_doc,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._load_docx,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": self._load_excel,
            "application/vnd.ms-excel": self._load_excel,
        }
        
        # Log available fallback processors
        if PDFPLUMBER_AVAILABLE:
            logger.info("pdfplumber available for PDF fallback processing")
        if PYTHON_DOCX_AVAILABLE:
            logger.info("python-docx available for DOCX fallback processing")
    
    def _calculate_content_hash(self, content: bytes) -> str:
        """Calculate SHA256 hash of document content for idempotency."""
        return hashlib.sha256(content).hexdigest()
    
    def _extract_document_metadata(self, file_path: str, content_type: str, content: bytes = None) -> Dict[str, Any]:
        """Extract enhanced metadata from document."""
        metadata = {
            "content_type": content_type,
            "file_size": len(content) if content else os.path.getsize(file_path),
            "processed_at": datetime.utcnow().isoformat(),
            "loader_version": "2.0"
        }
        
        # Add content hash for idempotency
        if content:
            metadata["content_hash"] = self._calculate_content_hash(content)
        
        # Extract file stats
        try:
            stat = os.stat(file_path)
            metadata.update({
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            })
        except Exception as e:
            logger.debug(f"Could not extract file stats: {e}")
        
        # Try to extract document-specific metadata
        try:
            if content_type == "application/pdf":
                metadata.update(self._extract_pdf_metadata(file_path))
        except Exception as e:
            logger.debug(f"Could not extract document-specific metadata: {e}")
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata."""
        metadata = {}
        try:
            doc = fitz.open(file_path)
            pdf_meta = doc.metadata
            if pdf_meta:
                metadata.update({
                    "title": pdf_meta.get("title", ""),
                    "author": pdf_meta.get("author", ""),
                    "subject": pdf_meta.get("subject", ""),
                    "creator": pdf_meta.get("creator", ""),
                    "producer": pdf_meta.get("producer", ""),
                    "creation_date": pdf_meta.get("creationDate", ""),
                    "modification_date": pdf_meta.get("modDate", "")
                })
            metadata["total_pages"] = len(doc)
            doc.close()
        except Exception as e:
            logger.debug(f"Failed to extract PDF metadata: {e}")
        return metadata
    
    def load_from_file(self, file_path: str, content_type: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """Load documents from file with enhanced metadata extraction."""
        try:
            if content_type not in self.supported_formats:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Extract enhanced metadata
            enhanced_metadata = self._extract_document_metadata(file_path, content_type)
            if metadata:
                enhanced_metadata.update(metadata)
            
            loader_func = self.supported_formats[content_type]
            documents = loader_func(file_path, enhanced_metadata)
            
            logger.info(f"Loaded {len(documents)} documents from {file_path} with enhanced metadata")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load document from {file_path}: {e}")
            raise
    
    def load_from_bytes(self, file_bytes: bytes, content_type: str, filename: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """Load documents from bytes with content hashing."""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=self._get_extension(content_type)) as tmp_file:
                tmp_file.write(file_bytes)
                tmp_file_path = tmp_file.name
            
            try:
                # Enhanced metadata with content hash
                file_metadata = metadata or {}
                file_metadata.update({
                    "filename": filename,
                    "content_type": content_type,
                    "size_bytes": len(file_bytes),
                    "content_hash": self._calculate_content_hash(file_bytes)
                })
                
                # Extract enhanced metadata
                enhanced_metadata = self._extract_document_metadata(tmp_file_path, content_type, file_bytes)
                enhanced_metadata.update(file_metadata)
                
                documents = self.load_from_file(tmp_file_path, content_type, enhanced_metadata)
                return documents
                
            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"Failed to load document from bytes: {e}")
            raise
    
    async def load_from_content(self, content: bytes, filename: str, content_type: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """Async wrapper for load_from_bytes."""
        return self.load_from_bytes(content, content_type, filename, metadata)
    
    def _load_pdf(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load PDF document using PyMuPDF with pdfplumber fallback."""
        try:
            return self._load_pdf_pymupdf(file_path, metadata)
        except Exception as e:
            logger.warning(f"PyMuPDF failed: {e}. Trying pdfplumber fallback...")
            if PDFPLUMBER_AVAILABLE:
                try:
                    return self._load_pdf_pdfplumber(file_path, metadata)
                except Exception as fallback_error:
                    logger.error(f"pdfplumber fallback also failed: {fallback_error}")
            
            # Ultimate fallback: try unstructured
            logger.warning("Attempting unstructured fallback for PDF")
            try:
                return self._load_with_unstructured(file_path, metadata)
            except Exception as final_error:
                logger.error(f"All PDF processing methods failed. Final error: {final_error}")
                raise e  # Raise original error
    
    def _load_pdf_pymupdf(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load PDF using PyMuPDF (primary method), combining all pages into single document."""
        pdf_doc = fitz.open(file_path)
        
        # Combine all pages into a single document for proper chunking
        all_pages_text = []
        page_numbers = []
        
        for page_num, page in enumerate(pdf_doc):
            text = page.get_text()
            if text.strip():  # Skip empty pages
                all_pages_text.append(text.strip())
                page_numbers.append(page_num + 1)
        
        pdf_doc.close()
        
        if not all_pages_text:
            logger.warning(f"No text content extracted from PDF: {file_path}")
            return []
        
        # Combine all page text with page separators
        combined_text = "\n\n".join(all_pages_text)
        
        # Create single document with combined text
        combined_metadata = metadata.copy()
        combined_metadata.update({
            "total_pages": len(pdf_doc),
            "pages_with_content": page_numbers,
            "extraction_method": "pymupdf"
        })
        
        return [Document(
            page_content=combined_text,
            metadata=combined_metadata
        )]
    
    def _load_pdf_pdfplumber(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load PDF using pdfplumber (fallback method), combining all pages into single document."""
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber not available")
        
        all_pages_text = []
        page_numbers = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    all_pages_text.append(text.strip())
                    page_numbers.append(page_num + 1)
        
        if not all_pages_text:
            logger.warning(f"No text content extracted from PDF: {file_path}")
            return []
        
        # Combine all page text with page separators
        combined_text = "\n\n".join(all_pages_text)
        
        # Create single document with combined text
        combined_metadata = metadata.copy()
        combined_metadata.update({
            "total_pages": len(page_numbers),
            "pages_with_content": page_numbers,
            "extraction_method": "pdfplumber"
        })
        
        return [Document(
            page_content=combined_text,
            metadata=combined_metadata
        )]
    
    def _load_text(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load text document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return [Document(
                page_content=text,
                metadata=metadata
            )]
            
        except Exception as e:
            logger.error(f"Failed to load text file: {e}")
            raise
    
    def _load_csv(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load CSV document with intelligent chunking for large files."""
        try:
            # First, check file size to decide on strategy
            file_size = os.path.getsize(file_path)
            
            # If file is large (> 1MB), use streaming approach
            if file_size > 1024 * 1024:  # 1MB threshold
                logger.info(f"Large CSV file detected ({file_size} bytes), using streaming approach")
                return self._load_csv_streaming(file_path, metadata)
            else:
                logger.info(f"Small CSV file ({file_size} bytes), using standard approach")
                return self._load_csv_standard(file_path, metadata)
                
        except Exception as e:
            logger.error(f"Failed to load CSV: {e}")
            raise
    
    def _load_csv_streaming(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load large CSV files in streaming chunks."""
        documents = []
        chunk_size = 100  # Process 100 rows at a time
        
        try:
            # Read CSV in chunks using pandas
            chunk_counter = 0
            
            for chunk in pd.read_csv(file_path, chunksize=chunk_size):
                chunk_counter += 1
                
                # Convert chunk to text representation
                chunk_text = self._convert_dataframe_to_text(chunk)
                
                if chunk_text.strip():
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        "chunk_number": chunk_counter,
                        "chunk_size": len(chunk),
                        "file_type": "csv",
                        "element_type": "table_chunk"
                    })
                    
                    documents.append(Document(
                        page_content=chunk_text,
                        metadata=chunk_metadata
                    ))
                
                # Log progress for large files
                if chunk_counter % 10 == 0:
                    logger.info(f"Processed {chunk_counter} CSV chunks ({chunk_counter * chunk_size} rows)")
            
            logger.info(f"Completed CSV streaming: {chunk_counter} chunks, {len(documents)} documents")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to stream CSV: {e}")
            # Fallback to standard approach with smaller chunks
            return self._load_csv_fallback(file_path, metadata)
    
    def _load_csv_standard(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load smaller CSV files using unstructured library."""
        try:
            # Use unstructured's partition function which will automatically detect CSV
            elements = partition(filename=file_path)
            
            documents = []
            for element in elements:
                if hasattr(element, 'text') and element.text.strip():
                    element_metadata = metadata.copy()
                    element_metadata.update({
                        "element_type": element.category if hasattr(element, 'category') else "table",
                        "file_type": "csv"
                    })
                    
                    documents.append(Document(
                        page_content=element.text,
                        metadata=element_metadata
                    ))
            
            # If no elements were extracted, try to read as plain text in chunks
            if not documents:
                logger.warning(f"No structured elements found in CSV, reading as chunked text")
                return self._load_csv_as_text_chunks(file_path, metadata)
            
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load CSV with unstructured: {e}")
            return self._load_csv_fallback(file_path, metadata)
    
    def _load_csv_fallback(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Fallback method to load CSV as text chunks."""
        try:
            logger.info("Using CSV fallback method with manual chunking")
            return self._load_csv_as_text_chunks(file_path, metadata)
        except Exception as e:
            logger.error(f"CSV fallback method failed: {e}")
            # Ultimate fallback - create a minimal document
            return [Document(
                page_content=f"Failed to load CSV file: {e}",
                metadata={**metadata, "error": True, "file_type": "csv"}
            )]
    
    def _load_csv_as_text_chunks(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load CSV as text with intelligent chunking."""
        documents = []
        max_chunk_size = 2000  # characters
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                # Read header first
                header = file.readline().strip()
                if not header:
                    return []
                
                current_chunk = header + "\n"
                chunk_counter = 1
                row_counter = 1
                
                for line in file:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check if adding this line would exceed chunk size
                    if len(current_chunk) + len(line) > max_chunk_size and len(current_chunk) > len(header):
                        # Save current chunk
                        chunk_metadata = metadata.copy()
                        chunk_metadata.update({
                            "chunk_number": chunk_counter,
                            "file_type": "csv",
                            "element_type": "text_chunk",
                            "rows_in_chunk": row_counter - 1
                        })
                        
                        documents.append(Document(
                            page_content=current_chunk.strip(),
                            metadata=chunk_metadata
                        ))
                        
                        # Start new chunk with header
                        current_chunk = header + "\n" + line + "\n"
                        chunk_counter += 1
                        row_counter = 2  # Header + this line
                    else:
                        current_chunk += line + "\n"
                        row_counter += 1
                
                # Don't forget the last chunk
                if current_chunk.strip() and len(current_chunk) > len(header):
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        "chunk_number": chunk_counter,
                        "file_type": "csv",
                        "element_type": "text_chunk",
                        "rows_in_chunk": row_counter - 1
                    })
                    
                    documents.append(Document(
                        page_content=current_chunk.strip(),
                        metadata=chunk_metadata
                    ))
            
            logger.info(f"CSV text chunking completed: {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load CSV as text chunks: {e}")
            raise
    
    def _convert_dataframe_to_text(self, df: 'pd.DataFrame') -> str:
        """Convert pandas DataFrame to readable text format."""
        try:
            # Get column names
            headers = list(df.columns)
            
            # Convert to string representation
            output = StringIO()
            
            # Write headers
            output.write(" | ".join(headers) + "\n")
            output.write("-" * (len(" | ".join(headers))) + "\n")
            
            # Write rows
            for _, row in df.iterrows():
                row_text = " | ".join([str(val) if pd.notna(val) else "" for val in row])
                output.write(row_text + "\n")
            
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to convert DataFrame to text: {e}")
            # Fallback to simple string representation
            return str(df.to_string(index=False))
    
    def _load_word_doc(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load legacy Word document (.doc) using unstructured."""
        try:
            return self._load_with_unstructured(file_path, metadata)
        except Exception as e:
            logger.error(f"Failed to load .doc file: {e}")
            raise
    
    def _load_docx(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load DOCX document with python-docx fallback."""
        try:
            return self._load_with_unstructured(file_path, metadata)
        except Exception as e:
            logger.warning(f"Unstructured failed for DOCX: {e}. Trying python-docx fallback...")
            if PYTHON_DOCX_AVAILABLE:
                try:
                    return self._load_docx_python_docx(file_path, metadata)
                except Exception as fallback_error:
                    logger.error(f"python-docx fallback also failed: {fallback_error}")
            raise e
    
    def _load_docx_python_docx(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load DOCX using python-docx (fallback method)."""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Skip empty rows
                        table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n" + "\n\n".join(tables_text)
            
            if all_text.strip():
                doc_metadata = metadata.copy()
                doc_metadata.update({
                    "extraction_method": "python_docx",
                    "paragraphs_count": len(paragraphs),
                    "tables_count": len(doc.tables)
                })
                
                return [Document(
                    page_content=all_text,
                    metadata=doc_metadata
                )]
            else:
                return []
            
        except Exception as e:
            logger.error(f"Failed to load DOCX with python-docx: {e}")
            raise
    
    def _load_with_unstructured(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load document using unstructured library, combining elements into a single document."""
        try:
            elements = partition(filename=file_path)
            
            # Combine all text elements into a single document for proper chunking
            text_parts = []
            element_types = []
            
            for element in elements:
                if hasattr(element, 'text') and element.text.strip():
                    text_parts.append(element.text.strip())
                    element_types.append(element.category if hasattr(element, 'category') else "unknown")
            
            if not text_parts:
                logger.warning(f"No text content extracted from {file_path}")
                return []
            
            # Combine all text with appropriate separators
            full_text = "\n\n".join(text_parts)
            
            # Create single document with combined text
            combined_metadata = metadata.copy()
            combined_metadata.update({
                "extraction_method": "unstructured",
                "element_count": len(text_parts),
                "element_types": list(set(element_types))
            })
            
            return [Document(
                page_content=full_text,
                metadata=combined_metadata
            )]
            
        except Exception as e:
            logger.error(f"Failed to load with unstructured: {e}")
            raise
    
    def _load_excel(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Load Excel document with sheet-by-sheet processing."""
        try:
            documents = []
            
            # Load Excel file with pandas
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            logger.info(f"Loading Excel file with {len(sheet_names)} sheets: {sheet_names}")
            
            for sheet_name in sheet_names:
                try:
                    # Read each sheet
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if df.empty:
                        logger.warning(f"Sheet '{sheet_name}' is empty, skipping")
                        continue
                    
                    # Convert sheet to text
                    sheet_text = self._convert_dataframe_to_text(df)
                    
                    if sheet_text.strip():
                        sheet_metadata = metadata.copy()
                        sheet_metadata.update({
                            "sheet_name": sheet_name,
                            "file_type": "excel",
                            "element_type": "spreadsheet",
                            "rows": len(df),
                            "columns": len(df.columns)
                        })
                        
                        documents.append(Document(
                            page_content=sheet_text,
                            metadata=sheet_metadata
                        ))
                        
                        logger.info(f"Processed Excel sheet '{sheet_name}': {len(df)} rows, {len(df.columns)} columns")
                
                except Exception as e:
                    logger.error(f"Failed to process Excel sheet '{sheet_name}': {e}")
                    # Create error document for this sheet
                    error_metadata = metadata.copy()
                    error_metadata.update({
                        "sheet_name": sheet_name,
                        "file_type": "excel",
                        "element_type": "error",
                        "error": str(e)
                    })
                    documents.append(Document(
                        page_content=f"Error processing Excel sheet '{sheet_name}': {e}",
                        metadata=error_metadata
                    ))
            
            if not documents:
                # Fallback: try to load as unstructured if no sheets processed
                logger.warning("No Excel sheets processed, trying unstructured fallback")
                return self._load_with_unstructured(file_path, metadata)
            
            logger.info(f"Successfully loaded Excel file: {len(documents)} sheets processed")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load Excel file: {e}")
            # Ultimate fallback: try unstructured
            try:
                logger.info("Attempting Excel fallback with unstructured")
                return self._load_with_unstructured(file_path, metadata)
            except Exception as fallback_error:
                logger.error(f"Excel fallback also failed: {fallback_error}")
                # Create minimal error document
                return [Document(
                    page_content=f"Failed to load Excel file: {e}",
                    metadata={**metadata, "error": True, "file_type": "excel"}
                )]

    def _get_extension(self, content_type: str) -> str:
        """Get file extension from content type."""
        extensions = {
            "application/pdf": ".pdf",
            "text/plain": ".txt",
            "text/markdown": ".md",
            "text/csv": ".csv",
            "application/csv": ".csv",
            "application/msword": ".doc",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
            "application/vnd.ms-excel": ".xls",
        }
        return extensions.get(content_type, ".txt")


# Global loader instance
document_loader = DocumentLoader()